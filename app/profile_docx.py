from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "templates" / "nvqs_profile_template.docx"
DOTS = ".............................................................................................................................................................................."


def generate_profile_docx(submission: dict[str, Any]) -> bytes:
    """Fill the official NVQS 09/GNN Word template.

    The template is an A3 landscape, two-column Word file converted once from
    the provided .doc. It intentionally has no Word tables, so filling is done
    by replacing the exact template paragraphs that contain each form line.
    """
    document = Document(str(TEMPLATE_PATH))
    payload = _payload(submission)
    personal = payload.get("personal_basic", {})
    family = payload.get("family_basic", {})
    marital = payload.get("marital_basic", {})

    _fill_profile_page(document, submission, personal, family, marital)
    _fill_family_page(document, payload, family, marital)
    _fill_self_page(document, payload, personal, submission)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _fill_profile_page(
    document: Document,
    submission: dict[str, Any],
    personal: dict[str, Any],
    family: dict[str, Any],
    marital: dict[str, Any],
) -> None:
    replacements = {
        "Họ, chữ đệm và tên khai sinh": "Họ, chữ đệm và tên khai sinh (viết chữ in hoa):\t{value}".format(
            value=_upper(submission.get("full_name"))
        ),
        "Họ, chữ đệm và tên thường dùng": "Họ, chữ đệm và tên thường dùng: {value}".format(
            value=_value(submission.get("full_name"))
        ),
        "Sinh ngày": "Sinh ngày {day} tháng {month} năm {year}    Giới tính (nam, nữ) {gender}".format(
            day=_date_part(personal.get("date_of_birth"), 0),
            month=_date_part(personal.get("date_of_birth"), 1),
            year=_date_part(personal.get("date_of_birth"), 2),
            gender=_value(personal.get("gender")),
        ),
        "Số thẻ căn cước/CCCD": f"Số thẻ căn cước/CCCD: {_value(submission.get('citizen_id_number'))}",
        "Nơi đăng ký khai sinh": f"Nơi đăng ký khai sinh:\t{_value(personal.get('birth_registration_place'))}",
        "Quê quán": f"Quê quán:\t{_value(personal.get('hometown'))}",
        "Dân tộc": "Dân tộc: {ethnicity}    Tôn giáo: {religion}    Quốc tịch: {nationality}".format(
            ethnicity=_value(personal.get("ethnicity")),
            religion=_value(personal.get("religion")),
            nationality=_value(personal.get("nationality")),
        ),
        "Nơi thường trú của gia đình": f"Nơi thường trú của gia đình:\t{_value(personal.get('family_permanent_residence'))}",
        "Nơi ở hiện tại của bản thân": "Nơi ở hiện tại của bản thân:\t{value}".format(
            value=_value(personal.get("current_residence") or personal.get("street_address"))
        ),
        "Thành phần gia đình": "Thành phần gia đình:\t{family_background}    Bản thân {personal_background}".format(
            family_background=_value(personal.get("family_background")),
            personal_background=_value(personal.get("personal_background")),
        ),
        "Trình độ giáo dục phổ thông": f"Trình độ giáo dục phổ thông: {_value(personal.get('education_level'))}",
        "Trình  độ đào tạo": "Trình độ đào tạo: {training_level}    Ngoại ngữ: {foreign_language}".format(
            training_level=_value(personal.get("training_level")),
            foreign_language=_value(personal.get("foreign_language")),
        ),
        "Chuyên ngành đào tạo": f"Chuyên ngành đào tạo: {_value(personal.get('major'))}",
        "Ngày vào Đảng CSVN": "Ngày vào Đảng CSVN:\t{party_join}    Chính thức: {party_official}".format(
            party_join=_value(personal.get("party_join_date")),
            party_official=_value(personal.get("party_official_date")),
        ),
        "Ngày vào Đoàn TNCS Hồ Chí Minh": f"Ngày vào Đoàn TNCS Hồ Chí Minh:\t{_value(personal.get('youth_union_join_date'))}",
        "Khen thưởng": "Khen thưởng:\t{reward}    Kỷ luật:\t{discipline}".format(
            reward=_value(personal.get("reward_record")),
            discipline=_value(personal.get("discipline_record")),
        ),
        "Nghề nghiệp": "Nghề nghiệp:\t{occupation}    Lương: Ngạch {salary_grade} bậc {salary_step}".format(
            occupation=_value(personal.get("occupation")),
            salary_grade=_value(personal.get("salary_grade")),
            salary_step=_value(personal.get("salary_step")),
        ),
        "Nơi làm việc": f"Nơi làm việc, (học tập):\t{_value(personal.get('workplace'))}",
        "Đã đi nước ngoài": "Đã đi nước ngoài (tên nước, thời gian, lý do): {value}".format(
            value=_value(personal.get("abroad_history"))
        ),
        "Họ tên cha": "Họ tên cha {name} (sống, chết) {status}".format(
            name=_value(family.get("father_name")),
            status=_value(family.get("father_status")),
        ),
        "Họ tên mẹ": "Họ tên mẹ: {name} (sống, chết) {status}".format(
            name=_value(family.get("mother_name")),
            status=_value(family.get("mother_status")),
        ),
        "Họ tên vợ": "Họ tên vợ (chồng):\t{name}    Sinh ngày {day} tháng {month} năm {year}".format(
            name=_value(marital.get("spouse_name")),
            day=_date_part(marital.get("spouse_date_of_birth"), 0),
            month=_date_part(marital.get("spouse_date_of_birth"), 1),
            year=_date_part(marital.get("spouse_date_of_birth"), 2),
        ),
        "Cha mẹ có": "Cha mẹ có {total} người con, {sons} trai {daughters} gái; bản thân là con thứ {order}".format(
            total=_value(family.get("total_children")),
            sons=_value(family.get("sons_count")),
            daughters=_value(family.get("daughters_count")),
            order=_value(family.get("birth_order")),
        ),
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for marker, replacement in replacements.items():
            if marker == "Sinh ngày" and "Giới tính" not in text:
                continue
            if marker == "Nghề nghiệp" and "Lương" not in text:
                continue
            if text.startswith(marker):
                _set_paragraph_text(paragraph, replacement)
                break

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Họ tên cha") and index + 1 < len(document.paragraphs):
            _set_paragraph_text(
                document.paragraphs[index + 1],
                "Sinh ngày {day} tháng {month} năm {year}. Nghề nghiệp:\t{occupation}".format(
                    day=_date_part(family.get("father_date_of_birth"), 0),
                    month=_date_part(family.get("father_date_of_birth"), 1),
                    year=_date_part(family.get("father_date_of_birth"), 2),
                    occupation=_value(family.get("father_occupation")),
                ),
            )
        if text.startswith("Họ tên mẹ") and index + 1 < len(document.paragraphs):
            _set_paragraph_text(
                document.paragraphs[index + 1],
                "Sinh ngày {day} tháng {month} năm {year}. Nghề nghiệp: {occupation}".format(
                    day=_date_part(family.get("mother_date_of_birth"), 0),
                    month=_date_part(family.get("mother_date_of_birth"), 1),
                    year=_date_part(family.get("mother_date_of_birth"), 2),
                    occupation=_value(family.get("mother_occupation")),
                ),
            )
        if text.startswith("Họ tên vợ") and index + 1 < len(document.paragraphs):
            _set_paragraph_text(
                document.paragraphs[index + 1],
                "Nghề nghiệp:\t{occupation}    Bản thân đã có {children_count} con".format(
                    occupation=_value(marital.get("spouse_occupation")),
                    children_count=_value(marital.get("children_count")),
                ),
            )


def _fill_family_page(document: Document, payload: dict[str, Any], family: dict[str, Any], marital: dict[str, Any]) -> None:
    lines = _family_lines(payload, family, marital)
    _fill_ruled_lines_after(document, "II. TÌNH HÌNH KINH TẾ, CHÍNH TRỊ CỦA GIA ĐÌNH", lines, stop_marker="III.")


def _fill_self_page(document: Document, payload: dict[str, Any], personal: dict[str, Any], submission: dict[str, Any]) -> None:
    lines = _self_lines(payload, personal)
    _fill_ruled_lines_after(document, "III. TÌNH HÌNH KINH TẾ, CHÍNH TRỊ, QUÁ TRÌNH CÔNG TÁC", lines, stop_marker="NGƯỜI KHAI")
    _replace_after_marker(document, "(Ký ghi rõ họ tên)", f"   (Ký ghi rõ họ tên)\n\n\n{_upper(submission.get('full_name'))}")


def _fill_ruled_lines_after(document: Document, start_marker: str, lines: list[str], *, stop_marker: str) -> None:
    in_range = False
    line_index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if start_marker in text:
            in_range = True
            continue
        if in_range and text.startswith(stop_marker):
            return
        if not in_range:
            continue
        if _is_ruled_line(text):
            next_text = lines[line_index] if line_index < len(lines) else DOTS
            _set_paragraph_text(paragraph, next_text)
            line_index += 1


def _family_lines(payload: dict[str, Any], family: dict[str, Any], marital: dict[str, Any]) -> list[str]:
    siblings = payload.get("siblings", []) if isinstance(payload.get("siblings"), list) else []
    children = payload.get("children", []) if isinstance(payload.get("children"), list) else []
    father_history = payload.get("father_history", []) if isinstance(payload.get("father_history"), list) else []
    mother_history = payload.get("mother_history", []) if isinstance(payload.get("mother_history"), list) else []
    lines = [
        "Cha: {name}; sinh năm {year}; nghề nghiệp {occupation}; tình trạng {status}; nơi ở hiện tại {residence}.".format(
            name=_value(family.get("father_name")),
            year=_value(_extract_year(family.get("father_date_of_birth"))),
            occupation=_value(family.get("father_occupation")),
            status=_value(family.get("father_status")),
            residence=_value(family.get("father_current_residence")),
        ),
        "Mẹ: {name}; sinh năm {year}; nghề nghiệp {occupation}; tình trạng {status}; nơi ở hiện tại {residence}.".format(
            name=_value(family.get("mother_name")),
            year=_value(_extract_year(family.get("mother_date_of_birth"))),
            occupation=_value(family.get("mother_occupation")),
            status=_value(family.get("mother_status")),
            residence=_value(family.get("mother_current_residence")),
        ),
    ]
    lines.extend(_history_lines("Quá trình của cha", father_history))
    lines.extend(_history_lines("Quá trình của mẹ", mother_history))
    if marital.get("spouse_name"):
        lines.append(
            "Vợ/chồng: {name}; sinh năm {year}; nghề nghiệp {occupation}; nơi ở hiện tại {residence}. {notes}".format(
                name=_value(marital.get("spouse_name")),
                year=_value(_extract_year(marital.get("spouse_date_of_birth"))),
                occupation=_value(marital.get("spouse_occupation")),
                residence=_value(marital.get("spouse_current_residence")),
                notes=_value(marital.get("spouse_notes"), empty=""),
            )
        )
    for index, sibling in enumerate(siblings, start=1):
        lines.append(
            "Anh/chị/em {index}: {name}; {relation}; sinh năm {year}; nghề nghiệp {occupation}; nơi học tập/làm việc {workplace}; nơi ở hiện tại {residence}. {notes}".format(
                index=index,
                name=_value(sibling.get("full_name")),
                relation=_value(sibling.get("relation")),
                year=_value(_extract_year(sibling.get("date_of_birth"))),
                occupation=_value(sibling.get("occupation")),
                workplace=_value(sibling.get("workplace")),
                residence=_value(sibling.get("current_residence")),
                notes=_value(sibling.get("notes"), empty=""),
            )
        )
    for index, child in enumerate(children, start=1):
        lines.append(
            "Con {index}: {name}; sinh năm {year}; học tập/nghề nghiệp {occupation}; nơi ở hiện tại {residence}. {notes}".format(
                index=index,
                name=_value(child.get("full_name")),
                year=_value(_extract_year(child.get("date_of_birth"))),
                occupation=_value(child.get("occupation")),
                residence=_value(child.get("current_residence")),
                notes=_value(child.get("notes"), empty=""),
            )
        )
    if family.get("family_notes"):
        lines.append(f"Ghi chú thêm về gia đình: {_value(family.get('family_notes'))}.")
    return _split_lines(lines)


def _self_lines(payload: dict[str, Any], personal: dict[str, Any]) -> list[str]:
    histories = payload.get("personal_history", []) if isinstance(payload.get("personal_history"), list) else []
    lines = []
    if histories:
        for item in histories:
            lines.append(
                "{from_year} - {to_year}: {stage}. {summary}".format(
                    from_year=_value(item.get("from_year")),
                    to_year=_value(item.get("to_year")),
                    stage=_value(item.get("stage_name")),
                    summary=_value(item.get("summary")),
                )
            )
    else:
        lines.append("Chưa có thông tin quá trình bản thân.")
    lines.append(f"Khen thưởng: {_value(personal.get('reward_record'))}.")
    lines.append(f"Kỷ luật: {_value(personal.get('discipline_record'))}.")
    return _split_lines(lines)


def _history_lines(title: str, items: list[dict[str, Any]]) -> list[str]:
    if not items:
        return [f"{title}: chưa có dữ liệu."]
    return [
        "{title}: {from_year} - {to_year}: {summary}.".format(
            title=title,
            from_year=_value(item.get("from_year")),
            to_year=_value(item.get("to_year")),
            summary=_value(item.get("summary")),
        )
        for item in items
    ]


def _split_lines(lines: list[str], *, max_length: int = 118) -> list[str]:
    output: list[str] = []
    for line in lines:
        text = str(line).strip()
        while len(text) > max_length:
            split_at = text.rfind(" ", 0, max_length)
            if split_at < 40:
                split_at = max_length
            output.append(text[:split_at].strip())
            text = text[split_at:].strip()
        if text:
            output.append(text)
    return output


def _replace_after_marker(document: Document, marker: str, replacement: str) -> bool:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(marker):
            _set_paragraph_text(paragraph, replacement)
            return True
    return False


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _is_ruled_line(text: str) -> bool:
    return bool(text) and text.count(".") >= 20 and not any(ch.isalpha() for ch in text)


def _payload(submission: dict[str, Any]) -> dict[str, Any]:
    payload = submission.get("payload", {})
    return payload if isinstance(payload, dict) else {}


def _date_part(value: Any, index: int) -> str:
    text = str(value or "").strip()
    if not text:
        return "......"
    separator = "/" if "/" in text else "-"
    parts = text.split(separator)
    if len(parts) == 3:
        if len(parts[0]) == 4:
            parts = [parts[2], parts[1], parts[0]]
        return parts[index] or "......"
    return text if index == 2 else "......"


def _extract_year(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    for separator in ("/", "-"):
        if separator in text:
            parts = text.split(separator)
            if len(parts) == 3:
                return parts[0] if len(parts[0]) == 4 else parts[-1]
    return text


def _upper(value: Any) -> str:
    return _value(value, empty="").upper()


def _value(value: Any, *, empty: str = "................................") -> str:
    text = str(value or "").strip()
    return text if text else empty
